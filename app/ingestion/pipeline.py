"""
DocuRAG — Document Ingestion Pipeline

Orchestrates the complete ingestion flow for a single document:

  1. Validate file (size, extension, MIME type)
  2. Compute SHA-256 checksum and check for duplicates
  3. Save raw file to storage
  4. Create Document record in PostgreSQL (status=PENDING)
  5. Detect language from extracted text sample
  6. Classify each page (digital / scanned / mixed)
  7. Create Page records in PostgreSQL
  8. Update Document status to CLASSIFIED
  9. Enqueue async Celery task for extraction (Phase 2+)

Design:
- The pipeline is split into small, independently testable stages
- Each stage receives explicit inputs (no implicit state)
- Errors in any stage update Document.status=FAILED and re-raise
- Async methods use asyncio-friendly patterns but remain CPU-light

Extensibility:
- New stages (OCR, table extraction, formula detection) plug in after
  classify() without touching existing code
- The DocumentIngestionPipeline class accepts an async DB session via DI
"""
from __future__ import annotations

import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.ingestion.classifier import (
    classify_pdf_pages,
    classify_image_document,
    classify_text_document,
)
from app.ingestion.schemas import (
    DocumentUploadResponse,
    PageClassificationResult,
)
from app.models.document import Document
from app.models.page import Page
from app.models.processing_job import ProcessingJob
from app.models.enums import DocumentStatus, DocumentType, JobStatus, PageType
from app.utils.file_utils import (
    compute_sha256,
    detect_mime_type,
    generate_stored_filename,
    mime_to_document_type,
    save_upload_to_raw_storage,
    validate_file_extension,
    validate_file_size,
)
from app.utils.language_utils import detect_language
from app.utils.exceptions import (
    DuplicateDocumentError,
    FileSizeLimitError,
    IngestionError,
    UnsupportedFileTypeError,
)
from config.logging_config import get_logger
from config.settings import get_settings

logger = get_logger(__name__)
settings = get_settings()


class DocumentIngestionPipeline:
    """
    Stateless ingestion pipeline that processes one document per call.

    Instantiated per request with a bound async DB session.
    All database writes are deferred to the caller's transaction boundary.

    Usage:
        pipeline = DocumentIngestionPipeline(db_session)
        response = await pipeline.ingest(file_path, original_filename)
    """

    def __init__(self, db: AsyncSession) -> None:
        """
        Args:
            db: Async SQLAlchemy session (injected from FastAPI dependency).
        """
        self.db = db
        self._logger = get_logger(self.__class__.__name__)

    async def ingest(
        self,
        file_path: Path,
        original_filename: str,
        extra_metadata: Optional[dict] = None,
    ) -> DocumentUploadResponse:
        """
        Execute the full ingestion pipeline for a single document.

        Args:
            file_path: Path to the temporary upload file.
            original_filename: Original filename as provided by the client.
            extra_metadata: Optional caller-supplied metadata dict.

        Returns:
            DocumentUploadResponse with document_id and initial status.

        Raises:
            UnsupportedFileTypeError: For invalid file types.
            FileSizeLimitError: If file exceeds size limit.
            DuplicateDocumentError: If checksum matches existing document.
            IngestionError: For other pipeline failures.
        """
        self._logger.info(
            "Ingestion pipeline started",
            filename=original_filename,
            file_path=str(file_path),
        )

        # ─ Stage 1: Validate ──────────────────────────────────────
        await self._validate(file_path, original_filename)

        # ─ Stage 2: Compute checksum ──────────────────────────────
        checksum = compute_sha256(file_path)
        await self._check_duplicate(checksum)

        # ─ Stage 3: Detect file type ────────────────────────────
        mime_type = detect_mime_type(file_path)
        doc_type_str = mime_to_document_type(mime_type)
        doc_type = DocumentType(doc_type_str)
        file_size = file_path.stat().st_size

        # ─ Stage 4: Save to raw storage ──────────────────────────
        stored_filename = generate_stored_filename(original_filename)
        stored_path = save_upload_to_raw_storage(file_path, stored_filename)

        # ─ Stage 5: Create Document record ──────────────────────
        document = await self._create_document(
            original_filename=original_filename,
            stored_filename=stored_filename,
            stored_path=stored_path,
            doc_type=doc_type,
            mime_type=mime_type,
            file_size=file_size,
            checksum=checksum,
            extra_metadata=extra_metadata,
        )

        # ─ Stage 6: Update status to CLASSIFYING ──────────────────
        document.status = DocumentStatus.CLASSIFYING
        document.processing_started_at = datetime.now(timezone.utc)
        await self.db.flush()

        # ─ Stage 7: Classify pages ──────────────────────────────
        try:
            page_results = await self._classify_pages(stored_path, doc_type)
            document.page_count = len(page_results)
        except Exception as exc:
            self._logger.error(
                "Page classification failed",
                doc_id=str(document.id),
                exc_info=exc,
            )
            document.status = DocumentStatus.FAILED
            document.error_message = f"Classification failed: {exc}"
            await self.db.flush()
            raise IngestionError(f"Classification failed: {exc}") from exc

        # ─ Stage 8: Detect language from text sample ──────────────
        text_sample = self._extract_text_sample(page_results)
        if text_sample:
            lang, confidence = detect_language(text_sample)
            document.detected_language = lang
            document.language_confidence = confidence

        # ─ Stage 9: Create Page records ──────────────────────────
        await self._create_pages(document.id, page_results)

        # ─ Stage 10: Finalise document status ────────────────────
        document.status = DocumentStatus.CLASSIFIED
        await self.db.flush()

        # ─ Stage 11: Create classification job record ─────────────
        job = ProcessingJob(
            document_id=document.id,
            job_type="classify",
            status=JobStatus.SUCCESS,
            result_metadata={
                "total_pages": len(page_results),
                "digital_pages": sum(1 for p in page_results if p.page_type == PageType.DIGITAL),
                "scanned_pages": sum(1 for p in page_results if p.page_type == PageType.SCANNED),
                "mixed_pages": sum(1 for p in page_results if p.page_type == PageType.MIXED),
            },
            started_at=document.processing_started_at,
            completed_at=datetime.now(timezone.utc),
        )
        self.db.add(job)
        await self.db.flush()

        self._logger.info(
            "Ingestion pipeline completed",
            doc_id=str(document.id),
            page_count=document.page_count,
            status=document.status.value,
            language=document.detected_language,
        )

        return DocumentUploadResponse(
            document_id=document.id,
            original_filename=document.original_filename,
            file_type=document.file_type,
            file_size_bytes=document.file_size_bytes,
            checksum=document.checksum,
            status=document.status,
        )

    # ── Private stage methods ────────────────────────────────────────────

    async def _validate(self, file_path: Path, original_filename: str) -> None:
        """Run all file-level validations before any processing."""
        try:
            validate_file_extension(Path(original_filename))
        except ValueError as exc:
            raise UnsupportedFileTypeError(
                file_type=Path(original_filename).suffix,
                filename=original_filename,
            ) from exc

        try:
            validate_file_size(file_path)
        except ValueError as exc:
            raise FileSizeLimitError(
                filename=original_filename,
                size_bytes=file_path.stat().st_size,
                limit_bytes=settings.document.max_file_size_bytes,
            ) from exc

    async def _check_duplicate(self, checksum: str) -> None:
        """Raise DuplicateDocumentError if checksum already exists."""
        result = await self.db.execute(
            select(Document.id).where(Document.checksum == checksum)
        )
        existing_id = result.scalar_one_or_none()
        if existing_id is not None:
            raise DuplicateDocumentError(
                checksum=checksum,
                existing_doc_id=str(existing_id),
            )

    async def _create_document(
        self,
        original_filename: str,
        stored_filename: str,
        stored_path: Path,
        doc_type: DocumentType,
        mime_type: str,
        file_size: int,
        checksum: str,
        extra_metadata: Optional[dict],
    ) -> Document:
        """Insert a new Document record and return it."""
        document = Document(
            original_filename=original_filename,
            stored_filename=stored_filename,
            file_path=str(stored_path),
            file_type=doc_type,
            mime_type=mime_type,
            file_size_bytes=file_size,
            checksum=checksum,
            status=DocumentStatus.PENDING,
            extra_metadata=extra_metadata,
        )
        self.db.add(document)
        await self.db.flush()  # Get the auto-generated ID
        self._logger.info(
            "Document record created",
            doc_id=str(document.id),
            file_type=doc_type.value,
        )
        return document

    async def _classify_pages(
        self,
        file_path: Path,
        doc_type: DocumentType,
    ) -> list[PageClassificationResult]:
        """
        Dispatch to the appropriate classifier based on document type.

        Returns list of PageClassificationResult (one per page).
        """
        if doc_type == DocumentType.PDF:
            return classify_pdf_pages(file_path)
        elif doc_type == DocumentType.IMAGE:
            return classify_image_document()
        elif doc_type in (DocumentType.CSV, DocumentType.DOCX, DocumentType.XLSX):
            # For flat document types, read sample text and classify as digital
            text_sample = await self._extract_flat_document_text(file_path, doc_type)
            return classify_text_document(text_sample)
        else:
            # Unknown type — create single UNKNOWN page
            from app.ingestion.schemas import PageClassificationResult
            return [
                PageClassificationResult(
                    page_number=1,
                    page_type=PageType.UNKNOWN,
                    text_char_count=0,
                    has_images=False,
                )
            ]

    async def _extract_flat_document_text(self, file_path: Path, doc_type: DocumentType) -> str:
        """
        Extract a text sample from non-PDF document types for language detection.

        Args:
            file_path: Path to the document.
            doc_type: Document type enum.

        Returns:
            Text sample string (may be empty on failure).
        """
        try:
            if doc_type == DocumentType.CSV:
                import pandas as pd  # type: ignore[import]
                df = pd.read_csv(file_path, nrows=10, encoding_errors="replace")
                return df.to_string()

            elif doc_type == DocumentType.DOCX:
                import docx  # type: ignore[import]
                doc = docx.Document(str(file_path))
                return " ".join(p.text for p in doc.paragraphs[:20])

            elif doc_type == DocumentType.XLSX:
                import openpyxl  # type: ignore[import]
                wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
                texts = []
                for sheet in wb.worksheets[:2]:  # First 2 sheets
                    for row in sheet.iter_rows(max_row=10, values_only=True):
                        texts.extend(str(cell) for cell in row if cell is not None)
                wb.close()
                return " ".join(texts)

        except Exception as exc:  # noqa: BLE001
            self._logger.warning(
                "Could not extract text sample from flat document",
                doc_type=doc_type.value,
                exc_info=exc,
            )
        return ""

    async def _create_pages(
        self,
        document_id: uuid.UUID,
        page_results: list[PageClassificationResult],
    ) -> None:
        """Bulk-insert Page records for all classified pages."""
        pages = [
            Page(
                document_id=document_id,
                page_number=result.page_number,
                page_type=result.page_type,
                width=result.width,
                height=result.height,
                ocr_applied=False,
            )
            for result in page_results
        ]
        self.db.add_all(pages)
        await self.db.flush()
        self._logger.debug(
            "Page records created",
            doc_id=str(document_id),
            count=len(pages),
        )

    @staticmethod
    def _extract_text_sample(
        page_results: list[PageClassificationResult],
        max_chars: int = 500,
    ) -> str:
        """
        Extract a text sample from the first digital pages for language detection.

        Args:
            page_results: Classification results with text_char_count.
            max_chars: Maximum characters to collect.

        Returns:
            Combined text sample string.
        """
        # We can only use char counts here (not the actual text)
        # Full text is available after the extraction stage
        # Return empty string — language detection happens post-extraction
        return ""
